from flask import Flask, request, send_file,jsonify
import re
import os
from io import BytesIO
import math
import spacy
from spacy.matcher import PhraseMatcher
import json
from flask_cors import CORS
import speech_recognition as sr
from docx import Document
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt

import sys
print("Starting app...", file=sys.stderr)


app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# Constants

nlp = spacy.load("en_core_web_sm")

TRANSPORT_FACTORS = {
    "personal": {
        "car": 0.12,
        "bike": 0.08,
        "electric_car": 0.04,
        "electric_scooter": 0.02,
        "bicycle": 0.0,
        "walk": 0.0
    },
    "public": {
        "bus": 0.06,
        "metro": 0.04,
        "train": 0.05,
        "diesel_train": 0.07,
        "auto": 0.09,
        "cab": 0.13,
        "e_rickshaw": 0.01
    },
    "air": {
        "flight_domestic": 0.18,
        "flight_international": 0.22
    }
}

SHOPPING_FACTORS = {
    "clothes": 2.0,
    "gadgets": 6.0,
    "groceries": 1.2
}

SHOPPING_COST_ESTIMATES = {
    "clothes": {"rupee_per_kg": 500},
    "gadgets": {"rupee_per_kg": 3000},
    "groceries": {"rupee_per_kg": 150}
}

WATER_FACTORS = {
    "tap": 0.25,
    "bottled": 1.5
}

PLASTIC_FACTORS = {
    "PET": 6.0,
    "HDPE": 4.0,
    "PVC": 5.0
}

FOOD_FACTORS = {
    "meat": 7.0,
    "beef": 60.0,
    "lamb": 24.0,
    "chicken": 6.0,
    "fish": 5.0,
    "egg": 4.5,
    "milk": 3.0,
    "paneer": 4.0,
    "cheese": 10.0,
    "yogurt": 2.2,
    "curd": 2.2,
    "butter": 12.0,
    "ghee": 9.0,
    "cream": 6.0,
    
    "rice": 2.5,
    "basmati": 2.4,
    "brown rice": 2.2,
    "wheat": 1.3,
    "flour": 1.2,
    "maida": 1.4,
    "bread": 1.8,
    "chapati": 1.3,
    "roti": 1.3,
    "paratha": 2.0,
    "poha": 1.2,
    "idli": 1.5,
    "dosa": 1.7,
    "upma": 1.6,

    "potato": 0.4,
    "onion": 0.3,
    "tomato": 0.4,
    "carrot": 0.3,
    "spinach": 0.2,
    "cabbage": 0.3,
    "cauliflower": 0.3,
    "brinjal": 0.4,
    "okra": 0.4,

    "apple": 0.4,
    "banana": 0.3,
    "orange": 0.5,
    "grapes": 0.6,
    "mango": 0.7,
    "pineapple": 0.7,
    "papaya": 0.5,
    "pomegranate": 0.6,

    "vegetables": 0.5,
    "fruits": 0.4,
    "vegan": 1.2,
    "junk": 7.0,
    "processed": 8.0,

    "chocolate": 19.0,
    "icecream": 3.5,
    "coffee": 17.0,
    "tea": 1.8,
    "sugar": 1.2,
    "oil": 6.0,
    "chips": 5.5,
    "biscuit": 3.2,
    "noodles": 4.0,
    "pizza": 6.0,
    "burger": 7.5,
    "soft drink": 3.0,
    "juice": 2.0
}


# Utility
def flatten_transport_factors():
    all_factors = {}
    for key, category in TRANSPORT_FACTORS.items():
        all_factors.update(category)
    return all_factors

def convert_to_standard(num, unit):
    if num is None:
        return 0
    try:
        value = float(num)
    except ValueError:
        return 0
    if not unit:
        return value
    unit = unit.lower()
    if unit in ["g", "gram", "grams"]:
        return value / 1000
    elif unit in ["ml", "milliliter", "milliliters"]:
        return value / 1000
    elif unit in ["kg", "kgs", "kilogram", "kilograms"]:
        return value
    elif unit in ["l", "liters", "litres", "liter"]:
        return value
    elif unit in ["km", "kilometers", "kilometres"]:
        return value
    elif unit in ["miles"]:
        return value * 1.60934
    return value

def is_overlapping(new_span, spans):
    for span in spans:
        if not (new_span[1] <= span[0] or new_span[0] >= span[1]):
            return True
    return False

base_dir = os.path.dirname(__file__)  # gets folder of app.py
file_path = os.path.join(base_dir, "aliases.json")
with open(file_path, encoding="utf-8") as f:
    raw_aliases = json.load(f)

phrase_matchers = {}

for category, mapping in raw_aliases.items():
    matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
    for standard, variants in mapping.items():
        matcher.add(standard, [nlp.make_doc(alias) for alias in variants])
    phrase_matchers[category] = matcher
    
def apply_phrase_matchers(user_input):
    doc = nlp(user_input)
    replaced_text = user_input

    for category, matcher in phrase_matchers.items():
        matches = matcher(doc)
        for match_id, start, end in matches:
            span = doc[start:end]
            label = nlp.vocab.strings[match_id]
            replaced_text = re.sub(rf"\b{re.escape(span.text)}\b", label, replaced_text, flags=re.IGNORECASE)

    return replaced_text

def detect_shopping_spacy(doc):
    shopping_keywords = {
        "clothes": [
            "shirt", "jeans", "clothes", "dress", "saree", "tshirt", "hoodie", "shoes",
            "jacket", "kurta", "suit", "trousers", "cap", "shorts"
        ],
        "gadgets": [
            "phone", "mobile", "laptop", "tablet", "camera", "charger", "tv", "watch",
            "headphones", "airpods", "smartwatch", "monitor", "mouse", "keyboard"
        ],
        "groceries": [
            "rice", "flour", "atta", "dal", "grocery", "groceries", "sugar", "salt",
            "oil", "bread", "wheat", "vegetables", "milk", "butter", "jam", 
            "apple", "banana", "soap", "detergent", "toothpaste", "snacks", 
            "coffee", "tea", "juice", "chocolate", "candy", "cake", "biscuits",
        ]
    }

    spend = 0
    category = None

    for token in doc:
        for cat, keywords in shopping_keywords.items():
            if token.lemma_ in keywords:
                category = cat

        if token.text.lower() in ["rs", "rupees"]:
            try:
                prev = token.nbor(-1)
                if prev.like_num:
                    spend = float(prev.text)
            except:
                pass

    return spend, category


def detect_food_spacy(doc):
    food_data = {}
    for token in doc:
        word = token.lemma_.lower()
        if word in FOOD_FACTORS:
            food_data[word] = food_data.get(word, 0) + 0.15  # Default to 150g
    return food_data


# Tips and Rewards
def generate_tips(result):
    tips = []

    transport = result.get("transport_total", 0)
    if transport > 100:
        tips.extend([
            f"Your transport emissions are very high ({transport} kg CO₂). Try limiting long-distance travel or combining errands to reduce trips.",
            "Switch to an electric or hybrid vehicle if possible.",
            "Use apps for ride-sharing or carpooling to reduce solo trips.",
            "Consider working remotely if your job allows to reduce commuting."
        ])
    elif transport > 50:
        tips.extend([
            f"Your transport emissions are quite high ({transport} kg CO₂). Use public transport more often.",
            "Try biking or walking for short distances.",
            "Plan your week to reduce unnecessary trips.",
        ])

    food = result.get("food_total", 0)
    if food > 70:
        tips.extend([
            f"Your food emissions are high ({food} kg CO₂). Reduce consumption of red meat and dairy.",
            "Buy local and seasonal produce to reduce transport-related emissions.",
            "Avoid food waste — freeze leftovers or plan meals in advance.",
            "Explore vegetarian or vegan recipes once or twice a week."
        ])
    elif food > 30:
        tips.extend([
            f"Your food-related emissions are {food} kg CO₂. Include more plant-based meals in your diet.",
            "Cut back on processed and packaged foods.",
            "Reduce portion sizes and compost food waste where possible.",
        ])

    plastic = result.get("plastic_kg", 0)
    if plastic > 5:
        tips.extend([
            f"You used {plastic} kg of plastic. Switch to glass, metal, or cloth alternatives.",
            "Carry a reusable bag, bottle, and straw when going out.",
            "Avoid products with excessive packaging.",
            "Buy in bulk to reduce plastic waste from packaging."
        ])
    elif plastic > 2:
        tips.extend([
            f"You used {plastic} kg of plastic. Try using shampoo bars and refillable containers.",
            "Opt for biodegradable packaging whenever possible.",
            "Participate in local plastic recycling or cleanup programs.",
        ])

    electricity = result.get("electricity_kwh", 0)
    if electricity > 50:
        tips.extend([
            f"You consumed {electricity} kWh of electricity. Switch to LED lights and unplug electronics when not in use.",
            "Use a programmable thermostat to optimize cooling/heating.",
            "Consider installing solar panels if you have the option.",
            "Wash clothes in cold water and air-dry them when possible."
        ])
    elif electricity > 20:
        tips.extend([
            f"Your electricity usage is {electricity} kWh. Reduce screen time and power-hungry devices.",
            "Turn off appliances at the socket instead of leaving them on standby.",
            "Use smart plugs to schedule or monitor appliance use.",
        ])

    shopping = result.get("shopping_spend", 0)
    if shopping > 50:
        tips.extend([
            f"Your shopping emissions are {shopping} kg CO₂. Avoid fast fashion; buy durable, timeless pieces.",
            "Support local businesses and eco-conscious brands.",
            "Think twice before buying: do I really need this?",
            "Repair and reuse items before replacing them."
        ])
    elif shopping > 10:
        tips.extend([
            f"Shopping contributed {shopping} kg CO₂. Try thrift stores and second-hand platforms.",
            "Limit impulse buys and unsubscribe from promotional emails.",
            "Choose digital or paperless alternatives where possible."
        ])

    water = result.get("water_liters", 0)
    if water > 200:
        tips.extend([
            f"You used {water} liters of water. Install low-flow showerheads and dual-flush toilets.",
            "Collect rainwater for gardening and cleaning.",
            "Run dishwashers and washing machines only with full loads.",
            "Avoid washing vehicles frequently or use waterless products."
        ])
    elif water > 100:
        tips.extend([
            f"Water consumption at {water} liters — shorten shower time to under 5 minutes.",
            "Turn off taps while brushing or shaving.",
            "Fix leaks promptly — even slow drips waste liters daily.",
        ])

    total = result.get("total_emission", 0)
    if total > 2000:
        tips.extend([
            f"Your total footprint is very high ({total} kg CO₂). A deep audit of your lifestyle could help — track and measure monthly emissions.",
            "Set goals to reduce 10% each month through small changes in every area.",
            "Get involved in community greening efforts or tree planting.",
        ])
    elif total > 1000:
        tips.extend([
            f"Total emissions of {total} kg CO₂ can be improved by tackling 2–3 habits at a time.",
            "Set personal sustainability challenges (like a no-buy month or plastic-free week).",
            "Educate others around you — collective change amplifies impact.",
        ])
    elif total < 200:
        tips.extend([
            "Great job! Your total carbon footprint is impressively low — you're on a sustainable path.",
            "Keep up your eco-friendly choices and inspire others by sharing your habits.",
            "Try offsetting what little CO₂ you emit via verified carbon offset platforms.",
        ])
    elif total < 100:
        tips.append("Excellent! You're living one of the lowest-impact lifestyles according to your total carbon footprint. You could mentor others on how to reduce theirs!")

    return tips

def assign_badges(result):
    badges = []

    if result["total_emission"] < 100:
        badges.append("Low Carbon Hero")
    
    if result["total_emission"] < 4000:
        badges.append("Below Global Average")

    if result.get("plastic_kg", 0) < 1:
        badges.append("Plastic Reducer")

    if result.get("transport_total", 0) < 10:
        badges.append("Eco Commuter")

    if result.get("food_total", 0) < 5:
        badges.append("Green Eater")

    if result.get("electricity_kwh", 0) < 10:
        badges.append("Energy Saver")

    if result.get("shopping_spend", 0) < 2:
        badges.append("Minimal Shopper")

    if result.get("water_liters", 0) < 10:
        badges.append("Water Wise")

    return badges

# Main Calculation
def calculate_carbon(
    transport_data={},            
    electricity_kwh=0,
    food_data={},       
    shopping_spend=0,
    shopping_type="clothes",
    flight_km=0,
    flight_type="domestic",
    water_liters=0,
    water_type="tap",
    plastic_kg=0,
    plastic_type="PET"
):
    result = {}
    total = 0
    unknowns = {"food": [], "plastic": []}

    # Transport
    all_transport = flatten_transport_factors()
    transport_emissions = 0
    transport_details = {}

    for mode, km in transport_data.items():
        factor = all_transport.get(mode.lower(), 0)
        emission = round(km * factor, 2)
        transport_details[mode] = emission
        transport_emissions += emission

    result["transport_details"] = transport_details
    result["transport_total"] = round(transport_emissions, 2)
    total += transport_emissions

    # Electricity
    elec_emission = round(electricity_kwh * 0.7, 2)
    result["electricity_kwh"] = elec_emission
    total += elec_emission

    # Food
    food_emissions = 0
    food_breakdown = {}
    for item, qty in food_data.items():
        item_lower = item.lower()
        factor = FOOD_FACTORS.get(item_lower, None)
        if factor is None:
            unknowns["food"].append(item)
            factor = 5.0
        emission = round(qty * factor, 2)
        food_breakdown[item] = emission
        food_emissions += emission

    result["food_details"] = food_breakdown
    result["food_total"] = round(food_emissions, 2)
    total += food_emissions

    # Shopping
    shop_data = SHOPPING_COST_ESTIMATES.get(shopping_type.lower(), {"rupee_per_kg": 100})
    shop_factor = SHOPPING_FACTORS.get(shopping_type.lower(), 1.5)
    estimated_kg = shopping_spend / shop_data["rupee_per_kg"]
    shop_emission = round(estimated_kg * shop_factor, 2)
    result["shopping_spend"] = shop_emission
    total += shop_emission

    # Flights
    if flight_km > 0 and flight_type:
        key = "flight_" + flight_type.lower()
        air_factor = TRANSPORT_FACTORS["air"].get(key, 0.18)
        flight_emission = round(flight_km * air_factor, 2)
        result["flight_km"] = flight_emission
        total += flight_emission

    # Water
    water_factor = WATER_FACTORS.get(water_type.lower(), 0.25)
    water_emission = round((water_liters / 100) * water_factor, 2)
    result["water_liters"] = water_emission
    total += water_emission

    # Plastic
    plastic_type = plastic_type.upper()
    plastic_factor = PLASTIC_FACTORS.get(plastic_type, None)
    if plastic_factor is None:
        unknowns["plastic"].append(plastic_type)
        plastic_factor = 5.0
    plastic_emission = round(plastic_kg * plastic_factor, 2)
    result["plastic_kg"] = plastic_emission
    total += plastic_emission

    # Category Breakdown
    if total > 0:
        result["category_percentages"] = {
            "transport": round(transport_emissions / total * 100, 1),
            "electricity": round(elec_emission / total * 100, 1),
            "food": round(food_emissions / total * 100, 1),
            "shopping": round(shop_emission / total * 100, 1),
            "flight": round(flight_emission / total * 100, 1) if flight_km else 0.0,
            "water": round(water_emission / total * 100, 1),
            "plastic": round(plastic_emission / total * 100, 1)
        }

    # Final Total
    result["total_emission"] = round(total, 2)
    result["unknown_inputs"] = unknowns
    
    trees_required = math.ceil(result['total_emission'] / 0.7)
    result["trees_required"] = trees_required
    result["tips"] = generate_tips(result)
    result["badges"] = assign_badges(result)
    
    return result

def extract_number(groups):
    for g in groups:
        if g and re.match(r"^\d+(\.\d+)?$", g):
            return float(g)
    return 0

def parse_input_to_data(user_input):
    user_input = user_input.lower()
    
    transport_data = {}
    food_data = {}
    electricity_kwh = 0
    shopping_spend = 0
    shopping_type = "clothes"
    flight_km = 0
    flight_type = "domestic"
    water_liters = 0
    water_type = "tap"
    plastic_kg = 0
    plastic_type = "PET"

    alias_applied = apply_phrase_matchers(user_input)

    # Extract matched labels from aliases
    doc = nlp(alias_applied)
    matched_items = set()
    for category, matcher in phrase_matchers.items():
        matches = matcher(doc)
        for match_id, start, end in matches:
            label = nlp.vocab.strings[match_id]
            matched_items.add(label.lower())

    user_input = alias_applied
    
    # --- TRANSPORT ---
    matched_transport_spans = []
    all_transport_modes = flatten_transport_factors().keys()

    for mode in all_transport_modes:
        patterns = [
            rf"{mode}.*?(\d+(\.\d+)?)\s*(km|kilometers?|miles)",
            rf"(\d+(\.\d+)?)\s*(km|kilometers?|miles).*?{mode}",
            rf"(rode|used|took|travelled|drove|drive|commuted|covered|went by|on a).*?{mode}.*?(\d+(\.\d+)?)\s*(km|kilometers?|miles)",
            rf"distance.*?(\d+(\.\d+)?)\s*(km|kilometers?|miles).*?{mode}",
            rf"i.*?{mode}.*?(\d+(\.\d+)?)\s*(km|kilometers?|miles)",
            rf"{mode}.*?(covered|went|ran|moved|trip|journey|ride|travelled).*?(\d+(\.\d+)?)\s*(km|kilometers?|miles)",
            rf"{mode}.*?(\d+(\.\d+)?)(km|kilometers?|miles).*?(ride|travel)?",
            rf"{mode}.*?(commuted|traveled|used).*?(\d+(\.\d+)?)\s*(km|kilometers?|miles)",
            rf"drove\s+(\d+(\.\d+)?)\s*(km|kilometers?|miles)\s+in\s+a\s+{mode}",
            rf"used\s+a\s+{mode}\s+for\s+(\d+(\.\d+)?)\s*(km|kilometers?|miles)"
        ]
        for pattern in patterns:
            for match in re.finditer(pattern, user_input):
                span = match.span()
                if is_overlapping(span, matched_transport_spans):
                    continue
                groups = match.groups()
                num = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
                unit = next((g for g in groups if g and g.lower() in ["km", "kilometers", "miles"]), None)
                if num and unit:
                    km = convert_to_standard(num, unit)
                    transport_data[mode] = transport_data.get(mode, 0) + km
                    matched_transport_spans.append(span)

    # --- ELECTRICITY ---
    matched_electricity_spans = []
    electricity_patterns = [
        r"(\d+(\.\d+)?)\s*(kwh|kilowatt-hours?)",
        r"(used|consumed).*?(\d+(\.\d+)?)\s*(kwh|kilowatt-hours?)",
        r"electricity.*?(\d+(\.\d+)?)\s*(kwh|kilowatt-hours?)",
        r"(\d+(\.\d+)?)\s*(kwh|kilowatt-hours?)\s*used"
    ]
    for pattern in electricity_patterns:
        for match in re.finditer(pattern, user_input):
            span = match.span()
            if is_overlapping(span, matched_electricity_spans):
                continue
            groups = match.groups()
            for group in groups:
                if group and re.match(r"\d+(\.\d+)?", group):
                    electricity_kwh += convert_to_standard(group, "kwh")
                    matched_electricity_spans.append(span)
                    break

    # --- FOOD ---
    matched_food_spans = []
    for item in FOOD_FACTORS:
        item = item.lower()
        patterns = [
            rf"(\d+(\.\d+)?)\s*(kg|kgs|g|gram|grams|ml|milliliters?|l|liters?|litres?)\s+(of\s+)?{item}\b",
            rf"{item}\s*(amount|weighed|weighing|measured|totaled)?\s*(is|was)?\s*(about|around)?\s*(\d+(\.\d+)?)\s*(kg|g|grams|ml|l|liters?|litres?)\b",
            rf"(ate|had|consumed)\s+(about|around)?\s*(\d+(\.\d+)?)\s*(kg|g|grams|ml|l|liters?)\s+(of\s+)?{item}\b",
            rf"\b{item}\b.*?(about|around)?\s*(\d+(\.\d+)?)\s*(kg|g|ml|l)",
            rf"{item}\s*[:\-]?\s*(\d+(\.\d+)?)\s*(kg|g|ml|l)"
        ]
        for pattern in patterns:
            for match in re.finditer(pattern, user_input):
                span = match.span()
                if is_overlapping(span, matched_food_spans):
                    continue
                groups = match.groups()
                num = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
                unit = next((g for g in groups if g and g.lower() in ["kg", "kgs", "g", "gram", "grams", "ml", "l", "liters", "litres", "milliliters"]), None)
                if num and unit:
                    value = convert_to_standard(num, unit)
                    food_data[item] = food_data.get(item, 0) + value
                    matched_food_spans.append(span)

    # --- SHOPPING ---
    matched_shopping_spans = []

    for item in SHOPPING_FACTORS.keys():
        # Rupee-based
        rupee_patterns = [
            rf"{item}.*?(for|cost|price|worth|at)?\s*₹?\s*(\d+(\.\d+)?)\s*(rs|rupees)?",
            rf"(bought|purchased|got|ordered).*?{item}.*?(for|cost|price|worth|at)?\s*₹?\s*(\d+(\.\d+)?)\s*(rs|rupees)?",
            rf"spent\s*₹?\s*(\d+(\.\d+)?)\s*(rs|rupees)?\s*(on)?\s*{item}",
            rf"₹?\s*(\d+(\.\d+)?)\s*(rs|rupees)?\s*(for|on)?\s*{item}",
            rf"purchase of\s*{item}.*?(cost|price|was)?\s*₹?\s*(\d+(\.\d+)?)\s*(rs|rupees)?",
            rf"my\s*{item}.*?(cost|price|was)?\s*₹?\s*(\d+(\.\d+)?)\s*(rs|rupees)?",
        ]
        for pattern in rupee_patterns:
            for match in re.finditer(pattern, user_input):
                span = match.span()
                if is_overlapping(span, matched_shopping_spans):
                    continue
                groups = match.groups()
                num = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
                if num:
                    shopping_spend += float(num)
                    shopping_type = item
                    matched_shopping_spans.append(span)

        # Weight-based
        weight_patterns = [
            rf"(bought|purchased|got)?\s*(\d+(\.\d+)?)\s*(kg|kgs|kilograms?)\s+(of\s+)?{item}",
            rf"{item}.*?(amount|weighed|weighing)?\s*(is|was)?\s*(\d+(\.\d+)?)\s*(kg|kgs|kilograms?)",
        ]
        for pattern in weight_patterns:
            for match in re.finditer(pattern, user_input):
                span = match.span()
                if is_overlapping(span, matched_shopping_spans):
                    continue
                groups = match.groups()
                num = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
                if num:
                    shopping_spend += float(num) * 100
                    shopping_type = item
                    matched_shopping_spans.append(span)

    if shopping_spend == 0:
        for label in matched_items:
            for cat in SHOPPING_FACTORS:
                if label in raw_aliases.get("shopping", {}).get(cat, []):
                    shopping_type = cat
                    shopping_spend = 2000  
                    break

    # --- FLIGHT ---
    matched_flight_spans = []
    flight_patterns = [
        r"(\d+(\.\d+)?)\s*(km|kilometers?|miles)\s+(domestic|international|business|economy)?\s*flight",
        r"(domestic|international|business|economy)?\s*flight\s+of\s+(\d+(\.\d+)?)\s*(km|kilometers?|miles)",
        r"flight.*?(\d+(\.\d+)?)\s*(km|kilometers?|miles).*?(domestic|international|business|economy)?"
    ]
    for pattern in flight_patterns:
        for match in re.finditer(pattern, user_input):
            span = match.span()
            if is_overlapping(span, matched_flight_spans):
                continue
            groups = match.groups()
            print("Flight match groups:", groups)
            distance = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
            unit = next((g for g in groups if g in ["km", "kilometers", "miles"]), None)
            ftype = next((g for g in groups if g in ["domestic", "international", "business", "economy"]), None)
            if ftype:
                flight_type = ftype
            if distance and unit:
                flight_km += convert_to_standard(distance, unit)
                matched_flight_spans.append(span)
                
    # --- WATER ---
    matched_water_spans = []
    water_patterns = [
        r"(\d+(\.\d+)?)\s*(liters?|litres?|l|ml).*?(tap|bottled)?\s*water",
        r"(tap|bottled)\s*water.*?(\d+(\.\d+)?)\s*(liters?|litres?|l|ml)",
        r"water\s*(tap|bottled).*?(\d+(\.\d+)?)\s*(liters?|litres?|l|ml)",
        r"drank\s*(\d+(\.\d+)?)\s*(liters?|litres?|l|ml)\s*of\s*(tap|bottled)?\s*water"
    ]
    for pattern in water_patterns:
        matched = False
        for match in re.finditer(pattern, user_input):
            groups = match.groups()
            num = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
            unit = next((g for g in groups if g and g.lower() in ["liters", "litres", "l", "ml"]), None)
            wtype = next((g for g in groups if g and g.lower() in ["tap", "bottled"]), None)
            if wtype:
                water_type = wtype.lower()
            if num and unit:
                water_liters += convert_to_standard(num, unit)
                matched = True
                break
        if matched:
            break
        
    # --- PLASTIC ---
    plastic_patterns = [
        r"used\s*(\d+(\.\d+)?)\s*(kg|g|gram|grams)\s*(of\s*)?(pet|hdpe|pvc)?\s*plastic",
        r"(\d+(\.\d+)?)\s*(kg|g|gram|grams)\s*(of\s*)?(pet|hdpe|pvc)?\s*plastic",
        r"plastic.*?(pet|hdpe|pvc)?\s*(\d+(\.\d+)?)\s*(kg|g|gram|grams)"
    ]
    matched = False
    for pattern in plastic_patterns:
        if matched:
            break
        for match in re.finditer(pattern, user_input):
            groups = match.groups()
            num = next((g for g in groups if g and re.match(r"\d+(\.\d+)?", g)), None)
            unit = next((g for g in groups if g and g.lower() in ["kg", "g", "gram", "grams"]), None)
            ptype = next((g for g in groups if g and g.lower() in ["pet", "hdpe", "pvc"]), None)
            if ptype:
                plastic_type = ptype.upper()
            if num and unit:
                plastic_kg += convert_to_standard(num, unit)
                matched = True
                break
        doc = nlp(user_input)

    # SPA_CY fallback: food
    if len(food_data) == 0:
        spacy_food_data = detect_food_spacy(doc)
        for item, qty in spacy_food_data.items():
            food_data[item] = food_data.get(item, 0) + qty

    # SPA_CY fallback: shopping
    if shopping_spend == 0:
        spacy_spend, spacy_category = detect_shopping_spacy(doc)
        if spacy_spend > 0:
            shopping_spend = spacy_spend
        if spacy_category:
            shopping_type = spacy_category

    #* Debug Print
    print("Transport Data:", transport_data)
    print("Electricity (kWh):", electricity_kwh)
    print("Food Data:", food_data)
    print("Shopping Spend:", shopping_spend)
    print("Flight KM:", flight_km)
    print("Flight Type:", flight_type)
    print("Water (liters):", water_liters)
    print("Plastic (kg):", plastic_kg)
    print("Plastic Type:", plastic_type)

    return calculate_carbon(
        transport_data=transport_data,
        electricity_kwh=electricity_kwh,
        food_data=food_data,
        shopping_spend=shopping_spend,
        shopping_type=shopping_type,
        flight_km=flight_km,
        flight_type=flight_type,
        water_liters=water_liters,
        water_type=water_type,
        plastic_kg=plastic_kg,
        plastic_type=plastic_type
    )
    
@app.route('/')
def home():
    print("Home route hit!", file=sys.stderr)
    return "Backend is working!"

@app.route('/api/calculate', methods=['POST'])
def api_calculate():
    try:
        data = request.get_json()
        if not data or 'user_input' not in data:
            return jsonify({'error': 'Invalid request format'}), 400
            
        result = parse_input_to_data(data['user_input'])
        
        # Return the result directly (not nested in 'data' property)
        return jsonify({
            'total_emission': result['total_emission'],
            'category_percentages': result.get('category_percentages', {}),
            'tips': result.get('tips', []),
            'badges': result.get('badges', []),  # 👈 ADD THIS
            'transport_total': result.get('transport_total', 0),
            'electricity_kwh': result.get('electricity_kwh', 0),
            'food_total': result.get('food_total', 0),
            'shopping_spend': result.get('shopping_spend', 0),
            'water_liters': result.get('water_liters', 0),
            'plastic_kg': result.get('plastic_kg', 0) 
        })
    except Exception as e:
        return jsonify({
            'error': str(e)
        }), 500

@app.route('/recognize', methods=['POST'])
def recognize_speech():
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file provided'}), 400

    audio_file = request.files['audio']
    file_path = "temp_audio.wav"
    audio_file.save(file_path)

    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(file_path) as source:
            audio = recognizer.record(source)

        text = recognizer.recognize_google(audio)
        return jsonify({'text': text})

    except sr.UnknownValueError:
        return jsonify({'error': 'Speech not understood'}), 400
    except sr.RequestError as e:
        return jsonify({'error': f'Request failed: {e}'}), 500
    except Exception as e:
        return jsonify({'error': f'Unexpected error: {str(e)}'}), 500
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

@app.route('/api/download-report', methods=['POST'])
def api_download_report():
    data = request.get_json()
    user_input = data.get('user_input', '').strip()
    
    if not user_input:
        return jsonify({'error': 'Empty input'}), 400

    try:
        result = parse_input_to_data(user_input)

        # --- Create Word Document ---
        doc = Document()
        doc.add_heading('🌱 Carbon Footprint Report', 0)

        # Summary
        doc.add_paragraph().add_run("Total Emissions: ").bold = True
        doc.add_paragraph(f"{result['total_emission']} kg CO₂", style='Intense Quote')

        doc.add_paragraph().add_run("📊 Emission Breakdown:").bold = True
        doc.add_paragraph(f"• Transport: {result['transport_total']} kg CO₂")
        doc.add_paragraph(f"• Electricity: {result['electricity_kwh']} kg CO₂")
        doc.add_paragraph(f"• Food: {result['food_total']} kg CO₂")
        doc.add_paragraph(f"• Shopping: {result['shopping_spend']} kg CO₂")
        if result.get('flight_km', 0) > 0:
            doc.add_paragraph(f"• Flight: {result['flight_km']} kg CO₂")
        doc.add_paragraph(f"• Water: {result['water_liters']} liters")
        doc.add_paragraph(f"• Plastic: {result['plastic_kg']} kg")
        
        doc.add_paragraph().add_run("🌳 Trees Required to Offset: ").bold = True
        doc.add_paragraph(f"{result['trees_required']} trees")
        
        doc.add_paragraph().add_run("🏅 Badges Earned: ").bold = True
        
        if result.get('badges'):
            for badge in result['badges']:
                p = doc.add_paragraph()
                run = p.add_run()
                print(badge)
                badge_image_path = f"badges/{badge}.png"
                
                if os.path.exists(badge_image_path):
                    run.add_picture(badge_image_path, width=Inches(0.4)) 
                    run.add_text(f"  {badge}") 
                else:
                    p.add_run(f"• {badge}") 
        else:
            doc.add_paragraph("• No badges earned yet!", style='List Bullet')
        
        global_avg_annual_kg = 4.8 * 1000
        global_avg_daily_kg = global_avg_annual_kg / 365
        total_emission_daily = result['total_emission'] 
        doc.add_paragraph().add_run("🌍 Global Average Comparison: ").bold = True   
        
        
        if total_emission_daily > global_avg_daily_kg:
            excess = total_emission_daily - global_avg_daily_kg
            percent = (excess / global_avg_daily_kg) * 100
            doc.add_paragraph(
                f"Your daily carbon footprint is {total_emission_daily:.2f} kg CO₂, "
                f"which is about {percent:.1f}% higher than the global average daily footprint "
                f"of 13.15 kg CO₂ (4.8 tons per year)."
            )
        elif total_emission_daily < global_avg_daily_kg:
            less = global_avg_daily_kg - total_emission_daily
            percent = (less / global_avg_daily_kg) * 100
            doc.add_paragraph(
                f"Great job! Your daily carbon footprint is {total_emission_daily:.2f} kg CO₂, "
                f"which is about {percent:.1f}% lower than the global average daily footprint "
                f"of 13.15 kg CO₂ (4.8 tons per year)."
            )
        else:
            doc.add_paragraph(
                f"Your daily carbon footprint matches the global average daily footprint "
                f"of 13.15 kg CO₂ (4.8 tons per year)."
            )

        # --- Create Pie Chart ---
        labels = ['Transport', 'Electricity', 'Food', 'Shopping', 'Plastic', 'Flight']
        values = [
            result.get('transport_total', 0),
            result.get('electricity_kwh', 0),
            result.get('food_total', 0),
            result.get('shopping_spend', 0),
            result.get('plastic_kg', 0),
            result.get('flight_km', 0)
        ]
        # Remove zero values to keep chart clean
        chart_data = [(l, v) for l, v in zip(labels, values) if v > 0]
        chart_labels, chart_values = zip(*chart_data)

        # Generate pie chart
        fig, ax = plt.subplots()
        ax.pie(chart_values, labels=chart_labels, autopct='%1.1f%%', startangle=140)
        ax.axis('equal')
        chart_buf = BytesIO()
        plt.savefig(chart_buf, format='png')
        plt.close(fig)
        chart_buf.seek(0)

        # Insert chart into Word doc
        doc.add_page_break()
        doc.add_heading('📈 Visual Breakdown', level=1)
        doc.add_picture(chart_buf, width=Inches(5.5))

        # Save the document to a buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name="carbon_report.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/download-tips', methods=['POST'])
def api_download_tips():
    data = request.get_json()
    user_input = data.get('user_input', '').strip()
    
    if not user_input:
        return jsonify({'error': 'Empty input'}), 400
    
    try:
        result = parse_input_to_data(user_input)
        tips = generate_tips(result)
        
        if not tips:
            return jsonify({'error': 'No tips available for the provided input'}), 404

        # Create Word document
        doc = Document()
        doc.add_heading('🌿 Personalized Carbon Reduction Tips', 0)

        intro = doc.add_paragraph()
        intro.add_run("Based on your input, here are some sustainability tips to reduce your carbon footprint:\n").italic = True

        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')

        doc.add_paragraph("\n💡 Small steps can lead to big impact. Stay green!", style='Intense Quote')

        # Save document to in-memory buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="carbon_tips.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
if __name__ == '__main__':
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))